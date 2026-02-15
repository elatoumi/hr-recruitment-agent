from typing import List, Dict, Optional
from pathlib import Path
from langchain_core.tools import tool
import re
import unicodedata
import os

# Lazy imports for optional dependencies
try:
    from transformers import AutoModel, AutoTokenizer
    import torch
    from PIL import Image
    DEEPSEEK_AVAILABLE = True
except ImportError:
    DEEPSEEK_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    from langchain_community.vectorstores import FAISS
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_core.documents import Document
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False
    # Define dummy Document class for type hinting if import fails
    class Document:
        pass


class RAGHandler:
    """Singleton handler for RAG operations."""
    _vectorstore = None
    _embeddings = None
    
    @classmethod
    def get_embeddings(cls):
        if cls._embeddings is None:
            # Use a lightweight embedding model
            cls._embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        return cls._embeddings
    
    @classmethod
    def get_vectorstore(cls):
        if cls._vectorstore is None and RAG_AVAILABLE:
            cls._vectorstore = FAISS.from_documents(
                [Document(page_content="Initial placeholder", metadata={"source": "system"})],
                cls.get_embeddings()
            )
        return cls._vectorstore
    
    @classmethod
    def add_document(cls, text: str, metadata: dict):
        if not RAG_AVAILABLE:
            return
            
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.create_documents([text], metadatas=[metadata])
        
        vs = cls.get_vectorstore()
        vs.add_documents(chunks)
        print(f"Added {len(chunks)} chunks to RAG for {metadata.get('name')}")

    @classmethod
    def search(cls, query: str, k: int = 4) -> List[Document]:
        if not RAG_AVAILABLE:
            return []
        vs = cls.get_vectorstore()
        return vs.similarity_search(query, k=k)


class DeepSeekOCRHandler:
    """Handler for DeepSeek-OCR model inference."""
    
    _model = None
    _tokenizer = None
    
    @classmethod
    def load_model(cls, model_name: str = 'deepseek-ai/DeepSeek-OCR'):
        if not DEEPSEEK_AVAILABLE:
            raise ImportError("Required libraries (transformers, torch, PIL) not installed.")
            
        if cls._model is None:
            print(f"Loading {model_name}...")
            # Use GPU if available, else CPU (inference will be slow on CPU)
            device = "cuda" if torch.cuda.is_available() else "cpu"
            dtype = torch.bfloat16 if device == "cuda" else torch.float32
            
            cls._tokenizer = AutoTokenizer.from_pretrained(
                model_name, 
                trust_remote_code=True
            )
            
            cls._model = AutoModel.from_pretrained(
                model_name,
                trust_remote_code=True,
                use_safetensors=True
            ).eval().to(device).to(dtype)
            
            print("DeepSeek-OCR Loaded.")
            
    @classmethod
    def process_image(cls, image_path: str) -> str:
        """Process a single image path and return text."""
        cls.load_model()
        
        prompt = "<image>\n<|grounding|>Convert the document to markdown. "
        
        # Determine strict arguments based on availability
        # The prompt suggests model.infer(...) method usage
        try:
            res = cls._model.infer(
                cls._tokenizer,
                prompt=prompt,
                image_file=image_path,
                base_size=1024,
                image_size=640,
                crop_mode=True,
                save_results=False, # We just want text
                test_compress=True
            )
            return res
        except Exception as e:
            return f"OCR Error: {str(e)}"

    @classmethod
    def process_pdf(cls, pdf_path: str) -> str:
        """Convert PDF to images and process each."""
        if not PDF2IMAGE_AVAILABLE:
            return "Error: pdf2image library not found. Cannot convert PDF to images for OCR."

        import tempfile
        temp_files = []
        try:
            images = convert_from_path(pdf_path)
            full_text = []

            for i, img in enumerate(images):
                # Write temp image to system temp dir (not uploads)
                fd, temp_img_path = tempfile.mkstemp(suffix=f"_page_{i}.jpg")
                os.close(fd)
                temp_files.append(temp_img_path)
                img.save(temp_img_path, "JPEG")

                page_text = cls.process_image(temp_img_path)
                full_text.append(f"--- Page {i+1} ---\n{page_text}")

            return "\n\n".join(full_text)

        except Exception as e:
            return f"PDF Processing Error: {str(e)}"
        finally:
            for tf in temp_files:
                try:
                    if os.path.exists(tf):
                        os.remove(tf)
                except OSError:
                    pass


@tool
def cv_parser_tool(file_path: str) -> dict:
    """
    Parse a CV/Resume file via DeepSeek-OCR (or fallback) AND index it into the RAG vector store.
    
    Supports: PDF, Images, DOCX, TXT.
    Automatically adds the extracted content to the shared FAISS index for semantic search.
    
    Returns a dictionary with 'content' and 'metadata'.
    """
    path = Path(file_path)
    if not path.exists():
        return {
            "success": False,
            "content": "",
            "metadata": {},
            "error": "File not found"
        }
        
    file_ext = path.suffix.lower()
    content = ""
    error = None

    
    try:
        if file_ext in ['.jpg', '.jpeg', '.png']:
            if DEEPSEEK_AVAILABLE:
                try:
                    content = DeepSeekOCRHandler.process_image(str(path))
                except Exception as ocr_err:
                    error = f"DeepSeek-OCR failed on image: {ocr_err}"
            else:
                error = "DeepSeek-OCR dependencies not installed; cannot extract text from images"
                
        elif file_ext == '.pdf':
            # Try DeepSeek-OCR first, then fall back to pypdf
            ocr_succeeded = False
            if DEEPSEEK_AVAILABLE and PDF2IMAGE_AVAILABLE:
                try:
                    content = DeepSeekOCRHandler.process_pdf(str(path))
                    if content and not content.startswith("Error:") and not content.startswith("PDF Processing Error:"):
                        ocr_succeeded = True
                except Exception as ocr_err:
                    error = f"DeepSeek-OCR failed, falling back to pypdf: {ocr_err}"

            if not ocr_succeeded:
                try:
                    import pypdf
                    reader = pypdf.PdfReader(file_path)
                    pages_text = [page.extract_text() or "" for page in reader.pages]
                    content = "\n".join(pages_text).strip()
                    if not content:
                        error = "pypdf extracted no text (scanned PDF?)"
                except ImportError:
                    content = path.read_text(encoding="utf-8", errors="ignore")
                    error = "Advanced PDF parsing unavailable (missing pypdf)"

        elif file_ext in ['.txt', '.md', '.py', '.json']:
            content = path.read_text(encoding="utf-8", errors="ignore")
            
        elif file_ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                 error = "python-docx not installed"
        
        else:
            content = path.read_text(encoding="utf-8", errors="ignore")

        # --- RAG INGESTION ---
        if content and RAG_AVAILABLE:
            metadata = {
                "source": str(path),
                "name": path.name,
                "format": file_ext,
                "size": path.stat().st_size
            }
            RAGHandler.add_document(content, metadata)
        # ---------------------

        return {
            "success": True,
            "content": content,
            "metadata": {
                "name": path.name,
                "format": file_ext,
                "size": path.stat().st_size
            },
            "error": error
        }

    except Exception as e:
        return {
            "success": False,
            "content": "",
            "metadata": {},
            "error": str(e)
        }


@tool
def text_cleaner_pipeline(text: str) -> str:
    """
    Clean and normalize extracted text.
    Removes emojis, normalizes unicode, and collapses whitespace.
    """
    if not text:
        return ""

    text = unicodedata.normalize("NFKD", text)
    text = text.encode("ascii", "ignore").decode("ascii")

    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"
        "\U0001F300-\U0001F5FF"
        "\U0001F680-\U0001F6FF"
        "\U0001F700-\U0001F77F"
        "\U0001F780-\U0001F7FF"
        "\U0001F800-\U0001F8FF"
        "\U0001F900-\U0001F9FF"
        "\U0001FA00-\U0001FAFF"
        "]+",
        flags=re.UNICODE,
    )
    text = emoji_pattern.sub("", text)

    text = re.sub(r"\s+", " ", text).strip()
    return text


@tool
def anonymizer_tool(text: str) -> dict:
    """
    Anonymize CV text to reduce hiring bias.

    Uses Regex to strip names, emails, and phone numbers.
    """
    if not text:
        return {
            "anonymized_text": "",
            "removed_entities": [],
            "entity_count": 0
        }

    removed = []

    # Patterns combined from PRs
    email_pattern = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
    phone_pattern = re.compile(r"(\+?\d{1,3}[\s\-]?)?(\(?\d{2,4}\)?[\s\-]?)?\d{3,4}[\s\-]?\d{3,4}")
    name_pattern = re.compile(r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+){1,2})\b")

    if email_pattern.search(text):
        removed.append("email")
        text = email_pattern.sub("[REDACTED_EMAIL]", text)

    if phone_pattern.search(text):
        removed.append("phone")
        text = phone_pattern.sub("[REDACTED_PHONE]", text)

    if name_pattern.search(text):
        removed.append("name")
        text = name_pattern.sub("[REDACTED_NAME]", text)

    return {
        "anonymized_text": text,
        "removed_entities": removed,
        "entity_count": len(removed)
    }


@tool
def batch_upload_handler(files: List[str]) -> List[Dict]:
    """
    Process multiple CV files in batch.
    Parses and cleans each file, returning a list of results.
    """
    results = []

    for file_path in files:
        parsed = cv_parser_tool.invoke({"file_path": file_path})
        if parsed.get("success"):
            parsed["content"] = text_cleaner_pipeline.invoke({"text": parsed["content"]})
        results.append(parsed)

    return results
